"""
Document Processor Module

Handles document loading, chunking, and preprocessing for RAG.
"""

import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes various document types (PDF, DOCX, TXT) and chunks them for RAG.

    Supports multiple chunking strategies and metadata extraction.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Target size for text chunks in characters
            chunk_overlap: Number of characters to overlap between chunks
            separators: List of separators for splitting (defaults to common ones)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

    async def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from file path.

        Args:
            file_path: Path to the document file

        Returns:
            Dict with 'content', 'metadata', and 'source'
        """
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            file_extension = path.suffix.lower()

            if file_extension == '.txt':
                content = await self._load_txt(file_path)
            elif file_extension == '.pdf':
                content = await self._load_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = await self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            metadata = {
                'filename': path.name,
                'file_type': file_extension,
                'file_size': path.stat().st_size
            }

            logger.info(f"Loaded document: {file_path} ({len(content)} chars)")

            return {
                'content': content,
                'metadata': metadata,
                'source': str(path)
            }

        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise

    async def _load_txt(self, file_path: str) -> str:
        """Load plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with latin-1 encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    async def _load_pdf(self, file_path: str) -> str:
        """Load PDF file using pypdf."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except ImportError:
            raise ImportError("pypdf not installed. Install with: pip install pypdf")
        except Exception as e:
            raise Exception(f"Failed to read PDF: {e}")

    async def _load_docx(self, file_path: str) -> str:
        """Load DOCX file using python-docx."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n\n".join(text_parts)

        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to read DOCX: {e}")

    async def chunk_document(
        self,
        document: Dict[str, Any],
        preserve_metadata: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Split document into chunks for embedding.

        Args:
            document: Document dict from load_document()
            preserve_metadata: Whether to include metadata in each chunk

        Returns:
            List of chunk dicts with 'content', 'metadata', 'source', 'chunk_index'
        """
        try:
            content = document['content']
            chunks = self._split_text(content)

            result = []
            for i, chunk_text in enumerate(chunks):
                chunk_dict = {
                    'content': chunk_text,
                    'source': document['source'],
                    'chunk_index': i
                }

                if preserve_metadata:
                    chunk_dict['metadata'] = {
                        **document['metadata'],
                        'total_chunks': len(chunks),
                        'chunk_index': i
                    }
                else:
                    chunk_dict['metadata'] = {}

                result.append(chunk_dict)

            logger.info(f"Split document into {len(result)} chunks")
            return result

        except Exception as e:
            logger.error(f"Failed to chunk document: {e}")
            raise

    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks using recursive character splitting.

        Args:
            text: Text to split

        Returns:
            List of text chunks
        """
        chunks = []
        current_chunk = ""

        # Clean and normalize text
        text = self._clean_text(text)

        # Split by separators in order of preference
        splits = self._split_text_recursive(text, self.separators)

        for split in splits:
            # If adding this split would exceed chunk size, save current chunk
            if len(current_chunk) + len(split) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())

                # Start new chunk with overlap from previous chunk
                if self.chunk_overlap > 0:
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + split
                else:
                    current_chunk = split
            else:
                current_chunk += split

        # Add remaining text
        if current_chunk:
            chunks.append(current_chunk.strip())

        return [chunk for chunk in chunks if chunk]

    def _split_text_recursive(
        self,
        text: str,
        separators: List[str]
    ) -> List[str]:
        """
        Recursively split text using separators.

        Args:
            text: Text to split
            separators: List of separators to try

        Returns:
            List of text splits
        """
        if not separators:
            return [text]

        separator = separators[0]
        remaining_separators = separators[1:]

        if separator:
            splits = text.split(separator)
            # Keep the separator in the splits
            splits = [s + separator for s in splits[:-1]] + [splits[-1]]
        else:
            splits = list(text)

        # Recursively split large chunks
        final_splits = []
        for split in splits:
            if len(split) > self.chunk_size and remaining_separators:
                final_splits.extend(
                    self._split_text_recursive(split, remaining_separators)
                )
            else:
                final_splits.append(split)

        return final_splits

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text.

        Args:
            text: Text to clean

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters that might interfere with processing
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        return text.strip()

    async def process_documents(
        self,
        file_paths: List[str]
    ) -> List[Dict[str, Any]]:
        """
        Load and chunk multiple documents.

        Args:
            file_paths: List of file paths to process

        Returns:
            List of all chunks from all documents
        """
        all_chunks = []

        for file_path in file_paths:
            try:
                # Load document
                document = await self.load_document(file_path)

                # Chunk document
                chunks = await self.chunk_document(document)

                all_chunks.extend(chunks)

            except Exception as e:
                logger.error(f"Failed to process document {file_path}: {e}")
                # Continue with other documents

        logger.info(f"Processed {len(file_paths)} documents into {len(all_chunks)} chunks")
        return all_chunks

    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a document without loading full content.

        Args:
            file_path: Path to the document

        Returns:
            Metadata dict
        """
        try:
            path = Path(file_path)

            metadata = {
                'filename': path.name,
                'file_type': path.suffix.lower(),
                'file_size': path.stat().st_size,
                'created_at': path.stat().st_ctime,
                'modified_at': path.stat().st_mtime
            }

            # Add document-specific metadata
            if path.suffix.lower() == '.pdf':
                metadata.update(await self._extract_pdf_metadata(file_path))
            elif path.suffix.lower() in ['.docx', '.doc']:
                metadata.update(await self._extract_docx_metadata(file_path))

            return metadata

        except Exception as e:
            logger.error(f"Failed to extract metadata from {file_path}: {e}")
            return {}

    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            return {
                'page_count': len(reader.pages),
                'pdf_metadata': reader.metadata if reader.metadata else {}
            }
        except Exception:
            return {}

    async def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata."""
        try:
            from docx import Document

            doc = Document(file_path)
            core_props = doc.core_properties

            return {
                'author': core_props.author if core_props.author else '',
                'title': core_props.title if core_props.title else '',
                'subject': core_props.subject if core_props.subject else '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }
        except Exception:
            return {}
